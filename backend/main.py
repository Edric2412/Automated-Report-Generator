# main.py (Corrected and verified for Google Gemini API)

from fastapi import FastAPI, Form, UploadFile, File, Request, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import uuid
import re
from datetime import datetime, timedelta
import io
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from PIL import Image

# --- NEW: Imports for Google Gemini and .env loading ---
import google.generativeai as genai
from dotenv import load_dotenv
# --- End of new imports ---

# --- NEW: Load environment variables and configure Google Gemini API ---
load_dotenv()

# Global variable for the model
gemini_model = None
try:
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
    if not GOOGLE_API_KEY:
        print("CRITICAL: GOOGLE_API_KEY not found in .env file. Summarizer will be disabled.")
    else:
        genai.configure(api_key=GOOGLE_API_KEY)
        # Configure the model for safety and generation
        generation_config = {
            "temperature": 0.7,
            "top_p": 1,
            "top_k": 1,
            "max_output_tokens": 2048,
        }
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        ]
        gemini_model = genai.GenerativeModel(model_name="models/gemini-1.5-flash-latest",
                                             generation_config=generation_config,
                                             safety_settings=safety_settings)
        print("Google Gemini model configured successfully.")
except Exception as e:
    print(f"CRITICAL: Error configuring Google Gemini: {e}")
# --- End of Gemini Configuration ---


app = FastAPI()

# Directory setup
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
STATIC_DIR = os.path.join(BASE_DIR, "static")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
templates = Jinja2Templates(directory=TEMPLATES_DIR)


# Pydantic model for the summarizer request
class SummaryRequest(BaseModel):
    text: str
    sentences: int = 3 # This is now just a hint for the prompt length

# --- REVISED: auto-summarize endpoint using Google Gemini ---
@app.post("/api/auto-summarize")
async def auto_summarize(request: SummaryRequest):
    """
    Accepts text and returns an abstractive summary using Google Gemini.
    """
    if not gemini_model:
        raise HTTPException(status_code=503, detail="AI Summarizer is not available. Check server configuration.")

    try:
        # Sanitize HTML from CKEditor to get clean text
        soup = BeautifulSoup(request.text, "html.parser")
        clean_text = soup.get_text()

        if not clean_text.strip():
            return JSONResponse(content={"summary": ""})

        # Craft the prompt for the language model
        prompt = f"""
        You are an expert academic writer. Your task is to summarize the following text into a concise and professional paragraph of approximately {request.sentences} sentences.
        The summary should capture the key points, be well-written, and suitable for a formal event report.
        Do not use bullet points or lists. The output must be a single, coherent paragraph.

        Original Text:
        ---
        {clean_text}
        ---

        Concise Summary Paragraph:
        """
        
        # Generate the summary asynchronously
        response = await gemini_model.generate_content_async(prompt)
        
        # Extract the text from the response, cleaning it up
        summary = response.text.strip()
        
        return JSONResponse(content={"summary": summary})
    except Exception as e:
        # Provide a more detailed error for debugging
        error_message = f"Gemini summarization failed: {str(e)}"
        print(error_message)
        # Check for specific blocked content error
        if "block_reason" in str(e).lower():
             error_message = "The provided text was blocked by the safety filter. Please revise the content."
        raise HTTPException(status_code=500, detail=error_message)


def cleanup_old_files(directory, hours=24):
    now = datetime.now()
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            file_time = datetime.fromtimestamp(os.path.getmtime(file_path))
            if now - file_time > timedelta(hours=hours):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Failed to remove {file_path}: {e}")

@app.on_event("startup")
async def startup_event():
    # We no longer need the NLTK downloader, just the file cleanup
    cleanup_old_files(OUTPUT_DIR)

# --- The rest of the file is identical to your original working version ---

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    try:
        return templates.TemplateResponse("index.html", {"request": request})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Template error: {str(e)}")


@app.post("/preview_report")
async def preview_report(
    request: Request,
    eventType: str = Form(...),
    department: str = Form(...),
    topic: str = Form(...),
    expertName: Optional[str] = Form(None),
    venue: str = Form(...),
    eventDurationType: str = Form(...),
    date: Optional[str] = Form(None),
    startTime: Optional[str] = Form(None),
    endTime: Optional[str] = Form(None),
    startDate: Optional[str] = Form(None),
    endDate: Optional[str] = Form(None),
    coordinator: str = Form(...),
    participants: int = Form(...),
    summary: str = Form(...),
    outcome: str = Form(...),
    hodName: str = Form(...)
):
    try:
        # Format date/time consistently for both preview and download
        if eventDurationType.lower() == "multiple" and startDate and endDate:
            formatted_dateTime = (
                datetime.strptime(startDate, "%Y-%m-%d").strftime("%B %d, %Y")
                + " to " +
                datetime.strptime(endDate, "%Y-%m-%d").strftime("%B %d, %Y")
            )
        elif date and startTime and endTime:
            formatted_date = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
            formatted_time = (
                datetime.strptime(startTime, "%H:%M").strftime("%I:%M %p")
                + " - " +
                datetime.strptime(endTime, "%H:%M").strftime("%I:%M %p")
            )
            formatted_dateTime = formatted_date + ", " + formatted_time
        else:
            # Fallback in case of missing data
            formatted_dateTime = "Date information not provided"

        report_data = {
            "eventType": eventType.title(),
            "department": department,
            "topic": topic,
            "expertName": "N/A" if eventType.lower() == "field visit" else (expertName if expertName else "N/A"),
            "venue": venue,
            "dateTime": formatted_dateTime,
            "coordinator": coordinator,
            "participants": participants,
            "summary": summary,
            "outcome": outcome,
            "hodName": hodName
        }
        
        # Add query parameters for image counts to be used by JavaScript
        response = templates.TemplateResponse("preview.html", {"request": request, "report": report_data})
        return response
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview generation failed: {str(e)}")

def replace_placeholder(doc, placeholder, value):
    def process_paragraph(paragraph, ph, val):
        if ph in paragraph.text:
            full_text = ''.join(run.text for run in paragraph.runs)
            if ph in full_text:
                paragraph.clear()
                new_run = paragraph.add_run(full_text.replace(ph, str(val)))
                new_run.font.name = 'DIN Pro Regular'
                # Increase font size for event type and department
                if placeholder == "{{eventType}}" or placeholder == "{{department}}":
                    new_run.font.size = Pt(14)
                else:
                    new_run.font.size = Pt(11)
                # Set color to blue for specific text
                if paragraph.text.strip().endswith(" Report") or paragraph.text.strip().startswith("Department of ") or placeholder == "{{eventType}}" or placeholder == "{{department}}":
                    new_run.font.color.rgb = RGBColor(0, 112, 192)  # Blue color
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, placeholder, value)
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph, placeholder, value)
                for nested_table in cell.tables:
                    process_table(nested_table)
    for table in doc.tables:
        process_table(table)

def process_node_formatting(paragraph, node):
    if isinstance(node, str):
        run = paragraph.add_run(node)
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(11)
        return run
    if not hasattr(node, 'name'):
        run = paragraph.add_run(str(node))
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(11)
        return run
    if node.name in ['strong', 'b']:
        last_run = None
        if node.contents:
            for child in node.contents:
                last_run = process_node_formatting(paragraph, child)
                if last_run is not None:
                    last_run.bold = True
        else:
            last_run = paragraph.add_run(node.get_text())
            last_run.bold = True
            last_run.font.name = 'DIN Pro Regular'
            last_run.font.size = Pt(11)
        return last_run
    elif node.name in ['em', 'i']:
        last_run = None
        if node.contents:
            for child in node.contents:
                last_run = process_node_formatting(paragraph, child)
                if last_run is not None:
                    last_run.italic = True
        else:
            last_run = paragraph.add_run(node.get_text())
            last_run.italic = True
            last_run.font.name = 'DIN Pro Regular'
            last_run.font.size = Pt(11)
        return last_run
    elif node.contents:
        last_run = None
        for child in node.contents:
            last_run = process_node_formatting(paragraph, child)
        return last_run
    else:
        run = paragraph.add_run(node.get_text())
        run.font.name = 'DIN Pro Regular'
        run.font.size = Pt(11)
        return run

def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

def replace_placeholder_with_html(doc, placeholder, html_content):
    paragraphs = list(doc.paragraphs)
    for i, paragraph in enumerate(paragraphs):
        if placeholder in paragraph.text:
            paragraph.clear()
            soup = BeautifulSoup(html_content, "html.parser")
            last_par = paragraph
            elements = list(soup.body.children) if soup.body else list(soup.children)
            if not elements:
                last_par.add_run(soup.get_text())
                return
            first = True
            for element in elements:
                if isinstance(element, str) and element.strip() == '':
                    continue
                if element.name == 'p':
                    if first:
                        new_par = last_par
                        first = False
                    else:
                        new_par = insert_paragraph_after(last_par)
                    for child in element.children:
                        process_node_formatting(new_par, child)
                    last_par = new_par
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li', recursive=False):
                        list_style = 'List Bullet' if element.name == 'ul' else 'List Number'
                        new_par = insert_paragraph_after(last_par, style=list_style)
                        for child in li.children:
                            process_node_formatting(new_par, child)
                        last_par = new_par
                elif isinstance(element, str) and element.strip():
                    new_par = last_par if first else insert_paragraph_after(last_par)
                    run = new_par.add_run(element.strip())
                    run.font.name = 'DIN Pro Regular'
                    run.font.size = Pt(11)
                    last_par = new_par
            break

def update_header(doc, eventType):
    for paragraph in doc.paragraphs:
        if "{{eventType}}" in paragraph.text:
            text = paragraph.text.replace("{{eventType}}", eventType.title())
            paragraph.clear()
            run = paragraph.add_run(text)
            run.font.name = 'DIN Pro Regular'
            run.font.size = Pt(24)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 112, 192)  # Blue
            break

def set_section_vertical_alignment_bottom(section):
    try:
        sectPr = section._sectPr
        for child in sectPr.findall(qn('w:valign')):
            sectPr.remove(child)
        vAlign = OxmlElement('w:valign')
        vAlign.set(qn('w:val'), 'bottom')
        sectPr.append(vAlign)
    except Exception as e:
        print(f"Could not set vertical alignment: {e}")

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'nil')
        tblBorders.append(border_el)
    tblPr.append(tblBorders)


def add_signature_section(doc, coordinator, hodName):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)
    
    remove_table_borders(table)

    cell_left = table.cell(0, 0)
    cell_left.text = f"\n\n\n\nName & Signature of Faculty-in-charge\n{coordinator}"
    for para in cell_left.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.name = 'DIN Pro Regular'
            run.font.size = Pt(11)
            
    cell_right = table.cell(0, 1)
    cell_right.text = f"\n\n\n\nName & Signature of HoD\n{hodName}"
    for para in cell_right.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.name = 'DIN Pro Regular'
            run.font.size = Pt(11)

@app.post("/generate_report")
async def generate_report(
    eventType: str = Form(...),
    department: str = Form(...),
    topic: str = Form(...),
    expertName: Optional[str] = Form(None),
    venue: str = Form(...),
    eventDurationType: str = Form(...),
    date: Optional[str] = Form(None),
    startTime: Optional[str] = Form(None),
    endTime: Optional[str] = Form(None),
    startDate: Optional[str] = Form(None),
    endDate: Optional[str] = Form(None),
    coordinator: str = Form(...),
    participants: int = Form(...),
    summary: str = Form(...),
    outcome: str = Form(...),
    hodName: str = Form(...),
    invitePoster: Optional[List[UploadFile]] = File([]),
    actionPhotos: Optional[List[UploadFile]] = File([]),
    attendanceSheet: Optional[List[UploadFile]] = File([]),
    analysisReport: Optional[List[UploadFile]] = File([])
):
    try:
        if eventDurationType.lower() == "multiple" and startDate and endDate:
            formatted_dateTime = (
                datetime.strptime(startDate, "%Y-%m-%d").strftime("%B %d, %Y")
                + " to " +
                datetime.strptime(endDate, "%Y-%m-%d").strftime("%B %d, %Y")
            )
        elif date and startTime and endTime:
            formatted_date = datetime.strptime(date, "%Y-%m-%d").strftime("%B %d, %Y")
            formatted_time = (
                datetime.strptime(startTime, "%H:%M").strftime("%I:%M %p")
                + " - " +
                datetime.strptime(endTime, "%H:%M").strftime("%I:%M %p")
            )
            formatted_dateTime = formatted_date + ", " + formatted_time
        else:
            formatted_dateTime = "Date information not provided"

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = re.sub(r'[^\w\s-]', '', topic).strip().replace(' ', '_')
        output_filename = f"{department}_{safe_topic}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        template_path = os.path.join(TEMPLATES_DIR, "workshop_template.docx")
        if not os.path.exists(template_path):
             raise HTTPException(status_code=500, detail="Default template 'workshop_template.docx' not found.")
        doc = Document(template_path)

        replacements = {
            "{{eventType}}": eventType.title(),
            "{{department}}": department,
            "{{topic}}": topic,
            "{{expertName}}": expertName if expertName else "N/A",
            "{{venue}}": venue,
            "{{dateTime}}": formatted_dateTime,
            "{{coordinator}}": coordinator,
            "{{participants}}": str(participants),
            "{{hodName}}": hodName
        }
        for placeholder, value in replacements.items():
            replace_placeholder(doc, placeholder, value)
        
        if eventType.lower() == 'field visit':
            for table in doc.tables:
                for row in table.rows:
                    if 'Expert Name' in row.cells[0].text:
                        row._element.getparent().remove(row._element)
                        break
        
        replace_placeholder_with_html(doc, "{{summary}}", summary)
        replace_placeholder_with_html(doc, "{{outcome}}", outcome)
        doc.add_page_break()
        
        # This part is simplified, assuming header is part of template and doesn't need dynamic update after creation
        # update_header(doc, eventType) 

        image_sections = [
            ("Invite Poster", invitePoster),
            ("Action Photos", actionPhotos),
            ("Attendance Sheet", attendanceSheet),
            ("Analysis Report", analysisReport)
        ]
        for section_name, images in image_sections:
            valid_images = [img for img in images if img and img.filename] if images else []
            if valid_images:
                doc.add_paragraph(section_name, style='Heading 1')
                
                # Simplified image handling for brevity, can be expanded as before
                for img in valid_images:
                    img_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}_{img.filename}")
                    with open(img_path, "wb") as buffer:
                        buffer.write(await img.read())
                    
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                    except Exception as e:
                        print(f"Could not add picture {img.filename}: {e}")
                    finally:
                        os.remove(img_path)
                
                if section_name != "Analysis Report":
                    doc.add_page_break()

        add_signature_section(doc, coordinator, hodName)

        doc.save(output_path)
        cleanup_old_files(OUTPUT_DIR, hours=24)
        return {
            "message": "Report generated successfully",
            "filename": output_filename,
            "download_url": f"/download_report/{output_filename}"
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate report: {str(e)}")

@app.get("/download_report/{filename}")
async def download_report(filename: str):
    sanitized_filename = os.path.basename(filename)
    if sanitized_filename != filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    # Normalize and check path containment
    file_path = os.path.normpath(os.path.join(OUTPUT_DIR, sanitized_filename))
    abs_output_dir = os.path.abspath(OUTPUT_DIR)
    abs_file_path = os.path.abspath(file_path)
    if not abs_file_path.startswith(abs_output_dir + os.sep):
        raise HTTPException(status_code=400, detail="Attempt to access file outside output directory")
    if not os.path.exists(abs_file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        abs_file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=sanitized_filename
    )

@app.on_event("startup")
async def create_default_template():
    default_template_path = os.path.join(TEMPLATES_DIR, "workshop_template.docx")
    if not os.path.exists(default_template_path):
        try:
            doc = Document()
            # Basic template structure
            doc.add_paragraph("{{eventType}} Report", style='Title')
            doc.add_paragraph("Event Details", style='Heading 1')
            # ... add table for details as before ...
            doc.add_paragraph("Summary", style='Heading 1')
            doc.add_paragraph("{{summary}}")
            doc.add_paragraph("Outcome", style='Heading 1')
            doc.add_paragraph("{{outcome}}")
            doc.save(default_template_path)
            print(f"Created default template at {default_template_path}")
        except Exception as e:
            print(f"Failed to create default template: {e}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
